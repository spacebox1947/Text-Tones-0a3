"""
# TEXT-TONEs for Python 3.6
#
# Generates text scores for musical interpretation!
# Woah!
#
"""
from string import punctuation
from random import randint
from docx import *
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from FileLoader import *
from DocumentParameters import *

#-----------------------------------------------------------
#----Global Variables---------------------------------------
#-----------------------------------------------------------

# Use weighted variables when finding music letters
USE_WEIGHTS = True
#USE_WEIGHTS = False
# Set all music letters to upper case
#UPCASE = False
UPCASE = True
music_list = ["a", "b", "c", "d", "e", "f", "g"]
weights = [25, 126, 170, 89, 18, 150, 130]
MAX = 256
for i in range(len(weights)):
    milk = weights[i] / MAX
    print("Chance of picking {}: {} in {}\t{}%".format(music_list[i], weights[i], MAX, milk))
# Directory and Source Files
DIR_PREFIX = "./Source Materials/"
FILES = ["Shotry.txt",
         "The Desert Drum by Robert Hichens-Edit.txt",
         "BB-14-Introduction.txt",
         "BB-14-EvaluationOfIndividualReports.txt",
         "BB-14-KnownsAndUnknowns.txt"]
global PART_COUNT

#-----------------------------------------------------------
#----Document Fucntions-------------------------------------
#-----------------------------------------------------------

def paraFormat(documentParameters, paragraph):
    """
    # Handles justification and linespacing.
    """
    paragraph_format = paragraph.paragraph_format
    #paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
    #paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_format.line_spacing = documentParameters.getLineSpacing()

def addsRun(documentParameters, paragraph, cRun, BOLD):
    """
    # When BOLD == True, formats the text for music;
    #     i.e.; bold, capitol, etc.
    # Else, it creates a run of the text.
    # can be customized for different colors and fonts for
    # two types of text.
    """
    run = paragraph.add_run(cRun)
    run.bold = BOLD
    font = run.font
    font.name = documentParameters.getFontName()
    if BOLD:
        font.size = Pt(documentParameters.getMusicPt())
        font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    else:
        font.size = Pt(documentParameters.getPt())
        font.color.rgb = RGBColor(0x90, 0x90, 0x90)

def getPartCountStr():
    global PART_COUNT
    if PART_COUNT < 10:
        return "0" + str(PART_COUNT)
    else:
        return str(PART_COUNT)

def getPartCount():
    global PART_COUNT
    return PART_COUNT

def incPartCount():
    global PART_COUNT
    PART_COUNT += 1

def setPartCount(newVal):
    global PART_COUNT
    PART_COUNT = newVal

def saveDocument(document, fileSaveName, randomizeSaveName, save):
    """
    # generate filename and save!
    """
    if randomizeSaveName:
        filestring = fileSaveName + str(randint(0, 2056)) + '.docx'
    else:
        filestring = fileSaveName + getPartCountStr() + '.docx'
        incPartCount()
    if save:
        document.save(filestring)
        print("Done! Saving file as... {}".format(filestring))
        

#-----------------------------------------------------------
#----Music Letter Fucntions---------------------------------
#-----------------------------------------------------------

def isMusic(letter):
    """
    # returns true or false
    # also, passes current letter and roll value to getWeights()
    """
    if letter in music_list:
        return True;
    else:
        return False

def getWeights(letter, val):
    """
    # Returns true or false based on weights
    """
    if val < weights[music_list.index(letter)]:
        return True
    else:
        return False

def makeRunArray(words):
    """
    # Returns two lists - 
    # musArray contains substrings of input words
    # boolArray contains T/F of different formatting of letters.
    #
    # ISSUE: Generates empty list values when USE_WEIGHTS = True
    """
    musArray = []
    boolArray = []
    s = ''
    c = 0
    m = None
    w = None
    b = None
    # initiliaze m, b, and w
    m = isMusic(words[c].lower())
    if USE_WEIGHTS and m:
        # set w, and b and m based on w
        w = getWeights(words[c].lower(), randint(0, 256))
        if m and w:
            b = m
        else:
            m = not m
            b = m
    else:
        b = m
    # now flip through the rest of the shit
    while(c < len(words)):
        """
        # Since there's no ctl val to set when a new letter is checked,
        # the state machine will flip and continue rolling until the
        # right value comes out of weight and isMusic, which is wrong.
        """
        # get the next m value
        m = isMusic(words[c].lower())
        if USE_WEIGHTS and m:
            # set w, and b and m based on w
            w = getWeights(words[c].lower(), randint(0, 256))
            if m != w:
                #print ("The weights are against it!")
                m = not m
        if b == m:
            s += words[c]
            c += 1
        elif b != m:
            musArray.append(s)
            boolArray.append(b)
            s = ''
            b = not b
    # append the last str and bool values            
    musArray.append(s)
    boolArray.append(b)
    return musArray, boolArray

#-----------------------------------------------------------
#----Miscellaneous String Functions-------------------------
#-----------------------------------------------------------
def stripPunctuation(word):
    """
    # Works on single strings and arrays.
    """
    if type(word) == type(list):
        for i in range(len(word)):
            word[i] = ''.join(s for s in word[i] if s not in punctuation)
        return word
    else:
        return ''.join(s for s in word if s not in punctuation)

def joinCurrentParagraph(newPWords):
    """
    # Turns our counted list of words in the current
    # paragraph into a string.
    # for efficiency's sake, of course.
    """
    return ' '.join(s for s in newPWords)

def makeFileStr(fileName):
    return DIR_PREFIX + fileName

#-----------------------------------------------------------

def asIsDocument(documentParameters, fileLoader, fileSaveName, randomizeSaveName, save):
    document = Document()
    if fileLoader.getSanitize():
        count = 0
        breaks_i = 0
        #print ("Breaks: {}".format(fileLoader.getBreaks()))
        while (breaks_i < len(fileLoader.getBreaks())):
            print ("Break [{}] of [{}].".format(breaks_i, len(fileLoader.getBreaks())-1))
            paragraph = document.add_paragraph()
            paraFormat(documentParameters, paragraph)
            # keep adding runs of whole lines until current index = a break
            while(count < fileLoader.getBreaks()[breaks_i]):
                txt, ctl = makeRunArray(fileLoader.getLine(count))
                for i in range(len(txt)):
                    if ctl[i] and UPCASE:
                        addsRun(documentParameters, paragraph, txt[i].upper(), ctl[i])
                    else:
                        addsRun(documentParameters, paragraph, txt[i], ctl[i])
                count += 1
            # breaks_i increments too many times with the following code.
            if count == fileLoader.getBreaks()[breaks_i]:
                breaks_i += 1
            if ((breaks_i + 1) < len(fileLoader.getBreaks()))\
               and (count < fileLoader.getLength()):
                #print ("Getting weird for breaks [{}] and count [{}]".format(breaks_i, count))
                txt, ctl = makeRunArray(fileLoader.getLine(count))
                for i in range(len(txt)):
                    if ctl[i] and UPCASE:
                        addsRun(documentParameters, paragraph, txt[i].upper(), ctl[i])
                    else:
                        addsRun(documentParameters, paragraph, txt[i], ctl[i])
                count += 1
    # try making a file without using breaks, cuz my code sucks
    else:
        count = 0
        paragraph = document.add_paragraph()
        paraFormat(documentParameters, paragraph)
        while count < fileLoader.getLength():
            if len(fileLoader.getLine(count)) >= 3:
                txt, ctl = makeRunArray(fileLoader.getLine(count))
                for i in range(len(txt)):
                    if ctl[i] and UPCASE:
                        addsRun(documentParameters, paragraph, txt[i].upper(), ctl[i])
                    else:
                        addsRun(documentParameters, paragraph, txt[i], ctl[i])
                count += 1
            else:
                paragraph = document.add_paragraph()
                paraFormat(documentParameters, paragraph)
                count += 1
    # generate filename and save! 
    saveDocument(document, fileSaveName, randomizeSaveName, save)
        
#-----------------------------------------------------------
print()
print("-----")
print()
print ("Available source files in {}...".format(DIR_PREFIX))
for i in FILES:
    print (makeFileStr(i))
print()
print("-----")
print()
# construct and print info about document parameters
desertParameters = DocumentParameters()
d = [16, 17, 1.8, True, 'Lucida Console']
desertParameters.documentParamsFromList(d)
print ("Testing DocumentParameters Class...")
print ("Using Parameters: {}".format(d))
desertParameters.repr()

# construct and file loader from a given text file
#desert_file = "./Source Materials/The Desert Drum by Robert Hichens-Edit.txt"
desert_file = "./Source Materials/Shotry.txt"
desert = FileLoader(desert_file, "Desert Drum", True)
desert.genFileArray()
print()
print("-----")
print()

# prints a lot of data from FileLoader class
#desert.printToConsole()
#----args:----asIsDocument(desertParameters, fileLoader, fileName, randomizeFileName, save?)
USE_WEIGHTS = True
#setPartCount(1)
#while getPartCount() < 8:
    #asIsDocument(desertParameters, desert, "Text-Tones_Shotry-Part-", False, True)
#asIsDocument(desertParameters, desert, "Text-Tone-Desert-AsIs-", True, False)
print()
print("-----")
print()

# 12 parts of BB14 Intro
alienParameters = DocumentParameters()
d = [14, 15, 1.5, True, 'Lucida Console']
alienParameters.documentParamsFromList(d)
print ("Testing DocumentParameters Class...")
print ("Using Parameters: {}".format(d))
alienParameters.repr()
alien_file = []
alien = []
for i in range(2, 5):
    alien_file = makeFileStr(FILES[i])
    print (makeFileStr(FILES[i]))
    #alien.append(FileLoader(alien_file, "PBB: Special Report 14", True))
    alien.append(FileLoader(alien_file, "PBB: Special Report 14", False))
    alien[i-2].genFileArray()
    print()

# Make Some 10 Part Scores 
USE_WEIGHTS = True
prefix = "Project Blue Book SR14-"
FILES[i][6:-4]
T_PARTS = 4
for i in range(len(alien)):
    setPartCount(1)
    if i == 1:
        T_PARTS = 4
    else:
        T_PARTS = 10
    saveName = prefix + FILES[i+2][6:-4] + "-PART-"
    print ("SaveName: {}{}".format(saveName, getPartCount()))
    while getPartCount() < T_PARTS+1:
        #asIsDocument(alienParameters, alien[i], saveName, False, False)
        asIsDocument(alienParameters, alien[i], saveName, False, True)
